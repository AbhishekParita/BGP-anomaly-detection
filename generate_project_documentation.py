"""
BGP Anomaly Detection Project Documentation Generator
Generates a comprehensive Word document explaining the entire project
Updated to include Drift Detection and Real-time Dashboard
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

class ProjectDocumentationGenerator:
    def __init__(self, output_path="BGP_Anomaly_Detection_Complete_Documentation.docx"):
        self.doc = Document()
        self.output_path = output_path
        self.setup_styles()
        
    def setup_styles(self):
        """Setup custom styles for the document"""
        # Title style
        styles = self.doc.styles
        
        # Main heading style
        if 'CustomHeading1' not in styles:
            heading1 = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1.font.size = Pt(18)
            heading1.font.bold = True
            heading1.font.color.rgb = RGBColor(0, 51, 102)
            
    def add_title_page(self):
        """Create an attractive title page"""
        title = self.doc.add_heading('BGP Anomaly Detection System', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = self.doc.add_paragraph('Comprehensive Technical Documentation')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.color.rgb = RGBColor(64, 64, 64)
        
        self.doc.add_paragraph()
        
        info = self.doc.add_paragraph('Real-time BGP Network Anomaly Detection using')
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.runs[0].font.size = Pt(12)
        
        tech = self.doc.add_paragraph('Machine Learning • Deep Learning • RPKI Validation')
        tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tech.runs[0].font.size = Pt(12)
        tech.runs[0].font.bold = True
        tech.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        
        self.doc.add_paragraph('\n\n\n')
        
        date_para = self.doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()
        
    def add_table_of_contents(self):
        """Add table of contents"""
        self.doc.add_heading('Table of Contents', 1)
        
        toc_items = [
            '1. Executive Summary',
            '2. System Architecture Overview',
            '3. Project Workflow',
            '4. Core Components',
            '   4.1 Data Collection (RIS Live WebSocket)',
            '   4.2 Feature Processing',
            '   4.3 Model Training',
            '   4.4 Multi-Model Detection',
            '   4.5 Ensemble Coordination',
            '   4.6 Correlation Engine',
            '   4.7 Drift Detection & Retraining',
            '5. Real-time Dashboard',
            '6. Database Architecture',
            '7. Drift Detection System',
            '   7.1 Drift Monitoring',
            '   7.2 Automatic Retraining',
            '   7.3 Hot-Swap Mechanism',
            '8. File-by-File Documentation',
            '9. BGP Features Explained',
            '10. Installation & Setup',
            '11. Usage Guide - Complete Commands',
            '12. Configuration Files',
            '13. Troubleshooting Guide',
            '14. Performance Metrics'
        ]
        
        for item in toc_items:
            p = self.doc.add_paragraph(item, style='List Number' if item[0].isdigit() else 'List Bullet')
            
        self.doc.add_page_break()
        
    def add_executive_summary(self):
        """Add executive summary section"""
        self.doc.add_heading('1. Executive Summary', 1)
        
        summary = """
This project implements a state-of-the-art real-time BGP (Border Gateway Protocol) anomaly 
detection system that combines multiple machine learning approaches with self-healing capabilities 
to identify and classify network security threats.

The system collects live BGP data from RIS Live WebSocket service and processes it through an 
8-service microservices architecture to detect various types of BGP anomalies including:
• Route Hijacks - Malicious advertisement of IP prefixes
• Route Leaks - Incorrect propagation of routing information
• BGP Flapping - Unstable routing causing network instability
• Peer Instability - Connection issues between BGP peers
• Unusual Path Patterns - Abnormal AS path lengths and structures

The detection mechanism employs a sophisticated three-component hybrid approach:
1. LSTM Autoencoder (Deep Learning) - Identifies temporal anomalies through reconstruction error
2. Isolation Forest (Machine Learning) - Detects outliers in feature space
3. Heuristic Rules (Domain Knowledge) - Applies deterministic thresholds based on BGP expertise

CRITICAL NEW ADDITION - Self-Healing System:
The system now includes automatic drift detection and model retraining capabilities. A dedicated 
drift monitor continuously tracks model performance and automatically triggers retraining when 
degradation is detected, ensuring sustained accuracy over time without manual intervention.
"""
        self.doc.add_paragraph(summary)
        
        self.doc.add_heading('Key Features:', 2)
        features = [
            'Real-time data collection from RIS Live WebSocket (5-10 minute latency)',
            'Ensemble machine learning with 3 independent detectors',
            '8-service microservices architecture for scalability',
            'PostgreSQL database with auto-cleanup (50,000 record limit per table)',
            'Live web dashboard with Chart.js visualization (10-second refresh)',
            'FastAPI for optimized REST API endpoints',
            'Configurable severity thresholds (Critical, High, Medium, Low)',
            'Comprehensive correlation engine for incident tracking',
            'AUTOMATIC drift detection with hourly monitoring',
            'AUTOMATIC model retraining with zero-downtime hot-swap',
            'Model backup system with version control',
            'Complete testing framework for validation'
        ]
        
        for feature in features:
            self.doc.add_paragraph(f'✓ {feature}', style='List Bullet')
        
        self.doc.add_heading('System Status (Current):', 2)
        status_items = [
            'Database: 50,000 records with real BGP data (30-day retention)',
            'Models: LSTM (647 KB), Isolation Forest (5.78 MB), Heuristic (0.98 KB)',
            'All 8 services: Operational and tested',
            'Drift Monitor: Active with hourly checks',
            'Retraining: 3 scripts tested and working (IF: 30s, Heuristic: 10s, LSTM: 5-10min)',
            'Dashboard: Live updates every 10 seconds',
            'API Response Time: < 1 second (optimized with LIMIT queries)',
            'Zero Downtime: Hot-swap mechanism verified'
        ]
        
        for status in status_items:
            self.doc.add_paragraph(f'• {status}', style='List Bullet')
            
        self.doc.add_page_break()
        
    def add_architecture_overview(self):
        """Add system architecture section"""
        self.doc.add_heading('2. System Architecture Overview', 1)
        
        arch_desc = """
The BGP Anomaly Detection System follows a modern microservices architecture with 8 independent 
services working together for real-time detection, correlation, and self-healing capabilities.
The system is designed for scalability, maintainability, and 24/7 autonomous operation.
"""
        self.doc.add_paragraph(arch_desc)
        
        self.doc.add_heading('8-Service Microservices Architecture:', 2)
        
        services = [
            ('Service 1: RIS Live Collector', 
             'services/ris_live_collector.py - WebSocket client connecting to RIS Live (rrc00.ripe.net). '
             'Collects real-time BGP data, buffers 50 messages, processes every 5 seconds. '
             'Stores raw data in bgp_raw_data table.'),
            
            ('Service 2: Feature Aggregator', 
             'services/feature_aggregator.py - Extracts 9 BGP features in 30-second windows. '
             'Computes announcements, withdrawals, total_updates, withdrawal_ratio, flap_count, '
             'path_length, unique_peers, message_rate, session_resets. Stores in bgp_features table.'),
            
            ('Service 3: LSTM Detector', 
             'services/lstm_detector.py - Deep learning temporal pattern detector using autoencoder. '
             'Uses 10-timestep sequences for reconstruction error analysis. '
             'Model: 647 KB, detects temporal anomalies. Stores in lstm_detections table.'),
            
            ('Service 4: Isolation Forest Detector', 
             'services/isolation_forest_detector.py - Statistical outlier detector using tree ensemble. '
             'Model: 5.78 MB with 200 estimators, detects feature space anomalies. '
             'Stores in if_detections table.'),
            
            ('Service 5: Heuristic Detector', 
             'services/heuristic_detector.py - Rule-based detector using domain knowledge thresholds. '
             'Model: 0.98 KB with 6 dynamic rules. Checks critical churn, severe path length, '
             'mass withdrawals. Stores in heuristic_detections table.'),
            
            ('Service 6: Ensemble Coordinator', 
             'services/ensemble_coordinator.py - Combines all 3 detector outputs using weighted voting. '
             'Computes final ensemble score and severity classification (CRITICAL/HIGH/MEDIUM/LOW). '
             'Stores in ensemble_results table.'),
            
            ('Service 7: Correlation Engine', 
             'services/correlation_engine.py - Groups related detections into incidents. '
             'Correlates by time windows, ASN, prefix patterns. Creates comprehensive incidents. '
             'Stores in correlated_incidents table.'),
            
            ('Service 8: Drift Monitor (NEW)', 
             'services/drift_monitor.py - Self-healing component that monitors model performance hourly. '
             'Detects drift by comparing 7-day baseline vs 24-hour current window. '
             'Automatically triggers retraining when degradation detected. '
             'Creates drift flags and reports in drift_flags and drift_reports tables.')
        ]
        
        for service_name, service_desc in services:
            self.doc.add_heading(service_name, 3)
            self.doc.add_paragraph(service_desc)
        
        self.doc.add_heading('Architecture Layers:', 2)
        
        layers = [
            ('Layer 1: Data Collection', 
             'RIS Live WebSocket service provides global BGP data from route collectors (rrc00-rrc26). '
             'Data latency: 5-10 minutes. Continuous streaming with automatic reconnection.'),
            
            ('Layer 2: Feature Processing', 
             'Real-time feature extraction from raw BGP messages using sliding windows. '
             '30-second aggregation interval with 5-second processing delay. Buffer size: 50 messages.'),
            
            ('Layer 3: Detection (3 Models)', 
             'Three independent detectors run in parallel:\n'
             '  • LSTM: Temporal pattern analysis (5-10min retraining)\n'
             '  • Isolation Forest: Statistical outliers (30s retraining)\n'
             '  • Heuristic: Rule-based detection (10s retraining)'),
            
            ('Layer 4: Ensemble & Correlation', 
             'Weighted voting combines detector outputs. Correlation engine groups related alerts. '
             'Final severity classification and incident creation.'),
            
            ('Layer 5: Storage (PostgreSQL)', 
             'Database: PostgreSQL 17.7 on localhost:5432\n'
             'Auto-cleanup: 50,000 record limit per table\n'
             'Retention: 30 days with triggers\n'
             '8 main tables: raw_data, features, 3 detections, ensemble, incidents, drift'),
            
            ('Layer 6: Presentation (Dashboard & API)', 
             'FastAPI: Optimized endpoints with LIMIT-based queries\n'
             'Dashboard: HTML/JavaScript with Chart.js, 10-second refresh\n'
             'Live visualizations: Time series, model breakdown, detection table'),
            
            ('Layer 7: Self-Healing (Drift System)', 
             'Hourly drift monitoring comparing baseline vs current performance\n'
             'Automatic retraining with hot-swap (zero downtime)\n'
             'Model versioning with backup system')
        ]
        
        for layer_name, layer_desc in layers:
            self.doc.add_heading(layer_name, 3)
            self.doc.add_paragraph(layer_desc)
            
        self.doc.add_page_break()
        
    def add_workflow_diagram(self):
        """Add detailed workflow explanation"""
        self.doc.add_heading('3. Project Workflow', 1)
        
        workflow = """
The system operates in two distinct phases: Training Phase (offline) and Detection Phase (real-time).
"""
        self.doc.add_paragraph(workflow)
        
        self.doc.add_heading('Phase 1: Training Phase (Offline)', 2)
        
        training_steps = [
            ('Step 1', 'Data Collection', 'Historical BGP data is loaded from CSV files containing 30+ days of network activity.'),
            ('Step 2', 'Feature Engineering', 'Extract 9 core BGP features: announcements, withdrawals, total_updates, withdrawal_ratio, flap_count, path_length, unique_peers, message_rate, session_resets.'),
            ('Step 3', 'Data Preprocessing', 'Apply StandardScaler normalization and create sliding window sequences (length=10) for temporal modeling.'),
            ('Step 4', 'Train LSTM Autoencoder', 'Train unsupervised LSTM model on "normal" traffic (filtered at 95th percentile). Threshold set at 95th percentile of reconstruction error.'),
            ('Step 5', 'Train Isolation Forest', 'Train tree-based anomaly detector with 200 estimators and 0.01 contamination factor.'),
            ('Step 6', 'Model Validation', 'Evaluate models on held-out test set, compute metrics (AUC-ROC, precision, recall).'),
            ('Step 7', 'Save Artifacts', 'Persist models, scalers, thresholds, and configuration to model_output/ directory.')
        ]
        
        for step_num, step_name, step_desc in training_steps:
            p = self.doc.add_paragraph()
            p.add_run(f'{step_num}: {step_name}').bold = True
            p.add_run(f'\n{step_desc}')
            
        self.doc.add_heading('Phase 2: Detection Phase (Real-time)', 2)
        
        detection_steps = [
            ('Step 1', 'Data Generation/Streaming', 'BGP BMP messages are generated (simulation) or received from network and streamed to Kafka topic "bgp-stream".'),
            ('Step 2', 'Message Consumption', 'hybrid_detector.py consumes messages from Kafka in real-time.'),
            ('Step 3', 'Feature Extraction', 'Extract same 9 features from each BGP message and maintain sliding window buffer.'),
            ('Step 4', 'ML Prediction', 'Pass feature sequence through LSTM Autoencoder and raw features through Isolation Forest.'),
            ('Step 5', 'Heuristic Evaluation', 'Apply rule-based checks for extreme values (critical churn, severe path length, mass withdrawal).'),
            ('Step 6', 'Ensemble Scoring', 'Combine scores using weighted Z-score normalization with optimized weights (IF: 0.5, LSTM: 0.5).'),
            ('Step 7', 'RPKI Validation', 'Query Routinator API to validate route origin against ROA database.'),
            ('Step 8', 'Severity Classification', 'Map ensemble score to severity level: CRITICAL (>2.54), HIGH (>2.14), MEDIUM (>1.82), LOW (>1.58), NORMAL.'),
            ('Step 9', 'Alert Generation', 'If anomaly detected, create alert record with details, scores, and RPKI status.'),
            ('Step 10', 'Persistence', 'Save alert to PostgreSQL database with timestamp and metadata.')
        ]
        
        for step_num, step_name, step_desc in detection_steps:
            p = self.doc.add_paragraph()
            p.add_run(f'{step_num}: {step_name}').bold = True
            p.add_run(f'\n{step_desc}')
            
        self.doc.add_page_break()
        
    def add_core_components_overview(self):
        """Add core components overview section"""
        intro = """
The BGP Anomaly Detection System consists of four major component categories that work 
together to provide comprehensive network security monitoring. Each component has specific 
responsibilities and interfaces with other components through well-defined protocols.
        """
        self.doc.add_paragraph(intro)
        
        # Component 1: Data Generation & Collection
        self.doc.add_heading('4.1 Data Generation & Collection Layer', 2)
        
        data_desc = """
This layer is responsible for creating, collecting, and storing BGP network data. It simulates 
realistic BGP traffic patterns and provides the foundation for all downstream processing.
        """
        self.doc.add_paragraph(data_desc)
        
        self.doc.add_heading('Key Modules:', 3)
        data_modules = [
            ('bmp_generator.py', 'Simulates BGP Monitoring Protocol (BMP) messages following RFC 7854. Generates multiple message types including Route Monitoring, Peer Up/Down, Statistics Reports, and supports complete BGP path attributes. Can simulate normal traffic and various anomaly patterns.'),
            ('stream_generator.py', 'Kafka producer that streams BGP data for real-time processing. Reads preprocessed datasets and publishes to the "bgp-stream" topic with configurable message rates.'),
            ('db_connector.py', 'Database abstraction layer providing CRUD operations for PostgreSQL. Manages connections, bulk insertions, temporal queries, and result persistence with proper error handling.')
        ]
        
        for module_name, module_desc in data_modules:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(f'{module_name}: ').bold = True
            p.add_run(module_desc)
            
        # Component 2: Model Training
        self.doc.add_heading('4.2 Model Training & Preprocessing Layer', 2)
        
        training_desc = """
This layer implements the offline machine learning pipeline that trains anomaly detection 
models on historical BGP data. It handles feature engineering, data preprocessing, model 
architecture design, training, and artifact persistence.
        """
        self.doc.add_paragraph(training_desc)
        
        self.doc.add_heading('Key Modules:', 3)
        training_modules = [
            ('bgp_lstm_pipeline.py', 'Complete LSTM Autoencoder training pipeline (672 lines). Implements unsupervised learning with sliding window sequences. Includes data loading, feature extraction for all 9 BGP metrics, StandardScaler normalization, temporal train/test split, model training with early stopping, threshold computation at 95th percentile, comprehensive visualization, and artifact saving (model, scaler, config).'),
            ('train_if_model.py', 'Isolation Forest training script. Trains ensemble of 200 decision trees with 0.01 contamination factor. Uses same preprocessing as LSTM for consistency. Saves model and feature list using joblib serialization.'),
            ('run_training.py', 'Entry point script for initiating LSTM training. Sets data paths and output directories, then invokes the main training pipeline.')
        ]
        
        for module_name, module_desc in training_modules:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(f'{module_name}: ').bold = True
            p.add_run(module_desc)
            
        self.doc.add_heading('Training Outputs:', 3)
        outputs = [
            'LSTM Model: lstm_best.h5 (Keras H5 format)',
            'Isolation Forest: isolation_forest.pkl (Joblib pickle)',
            'StandardScaler: scaler.pkl (fitted on training data)',
            'Configuration: config.json (hyperparameters, thresholds)',
            'Evaluation Metrics: evaluation.json (AUC-ROC, precision, recall)',
            'Training Plots: Loss curves, error distributions, threshold analysis'
        ]
        
        for output in outputs:
            self.doc.add_paragraph(output, style='List Bullet')
            
        # Component 3: Ensemble Detection
        self.doc.add_heading('4.3 Ensemble Detection & Scoring Engine', 2)
        
        ensemble_desc = """
This is the core intelligence layer that combines multiple detection approaches into a unified 
scoring system. It loads pre-trained models, processes BGP features, computes anomaly scores 
from multiple sources, and fuses them using weighted Z-score normalization.
        """
        self.doc.add_paragraph(ensemble_desc)
        
        self.doc.add_heading('Key Module:', 3)
        
        ensemble_detail = """
ensemble_bgp_optimized.py (867 lines) - The central scoring engine implementing:
        """
        self.doc.add_paragraph(ensemble_detail)
        
        ensemble_features = [
            'HeuristicDetector Class: Implements deterministic rules based on domain expertise. Checks for critical churn (>2000 updates), severe path length (>25 hops), and mass withdrawals (>90% ratio).',
            
            'Model Loading: Loads Isolation Forest, LSTM Autoencoder, StandardScaler, and configuration from model_output/ directory.',
            
            'Isolation Forest Scoring: Computes anomaly scores using score_samples() method. Higher scores indicate stronger outliers in feature space.',
            
            'LSTM Scoring: Creates sliding window sequences, passes through autoencoder, computes reconstruction error. High error indicates deviation from normal patterns.',
            
            'Z-Score Normalization: Standardizes scores from both models to comparable scales (mean=0, std=1).',
            
            'Weighted Fusion: Combines normalized scores using optimized weights (IF: 0.5, LSTM: 0.5). Formula: ensemble_score = w_if × z_if + w_lstm × z_lstm',
            
            'Severity Mapping: Classifies ensemble scores into severity levels using thresholds from config:\n  • CRITICAL: score > 2.54 (top 0.5%)\n  • HIGH: score > 2.14 (top 2%)\n  • MEDIUM: score > 1.82 (top 5%)\n  • LOW: score > 1.58 (top 10%)\n  • NORMAL: score ≤ 1.58',
            
            'Ensemble Optimization: Grid search function to find optimal model weights that maximize detection accuracy while minimizing false positives.',
            
            'Visualization Generation: Creates comprehensive plots including score distributions, time series, threshold analysis, and model agreement charts.'
        ]
        
        for feature in ensemble_features:
            self.doc.add_paragraph(f'• {feature}', style='List Bullet')
            
        # Component 4: Real-time Monitoring
        self.doc.add_heading('4.4 Real-time Monitoring & Alert System', 2)
        
        monitoring_desc = """
This layer implements the production deployment system that consumes live BGP data streams, 
performs real-time inference, validates routes cryptographically, and generates alerts for 
security operations teams.
        """
        self.doc.add_paragraph(monitoring_desc)
        
        self.doc.add_heading('Key Modules:', 3)
        monitoring_modules = [
            ('hybrid_detector.py', 'Main real-time detector (272 lines). Consumes Kafka messages from "bgp-stream" topic, maintains sliding window buffer of 10 records for LSTM, performs ML prediction, queries Routinator API for RPKI/ROA validation, classifies severity, and persists alerts to PostgreSQL using SQLAlchemy engine.'),
            
            ('anomaly_pipeline.py', 'Orchestration script for batch processing (109 lines). Initializes simulator with router and peer configurations, generates structured BGP data, inserts to database, fetches time-windowed data for ML processing, invokes ensemble scoring, persists results, and reports summary statistics.')
        ]
        
        for module_name, module_desc in monitoring_modules:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(f'{module_name}: ').bold = True
            p.add_run(module_desc)
            
        self.doc.add_heading('Integration Points:', 3)
        integrations = [
            'Kafka Integration: Subscribes to "bgp-stream" topic with JSON deserialization, auto-commit enabled, starting from latest offset.',
            
            'RPKI Validation: Async HTTP calls to Routinator API (http://127.0.0.1:3323/api/v1/validity/{asn}/{prefix}). Returns validation status: valid, invalid, or not-found.',
            
            'Database Persistence: Uses SQLAlchemy engine to write alerts with full context including timestamp, ASN, prefix, all feature values, ML scores, RPKI status, and severity classification.',
            
            'Environment Configuration: Loads database credentials from environment variables using python-dotenv, supporting multiple deployment environments.'
        ]
        
        for integration in integrations:
            self.doc.add_paragraph(f'• {integration}', style='List Bullet')
            
        # Data Flow Summary
        self.doc.add_heading('Component Interaction Flow', 2)
        
        flow_desc = """
The components interact in a well-defined sequence:
        """
        self.doc.add_paragraph(flow_desc)
        
        flow_steps = [
            '1. Data Generation Layer produces/receives BGP messages',
            '2. Messages are streamed through Kafka message queue',
            '3. Real-time Monitoring Layer consumes messages',
            '4. Features are extracted and buffered',
            '5. Ensemble Detection Layer scores the data using pre-trained models',
            '6. RPKI validation checks route origin authenticity',
            '7. Severity classification determines alert level',
            '8. Alerts are persisted to database for analysis',
            '9. Visualization tools generate reports and dashboards'
        ]
        
        for step in flow_steps:
            self.doc.add_paragraph(step, style='List Bullet')
    
    def add_bgp_features_section(self):
        """Add detailed BGP features explanation"""
        self.doc.add_heading('6. BGP Features Explained', 1)
        
        intro = """
The system extracts and analyzes 9 critical BGP features that capture different aspects 
of network behavior. These features are carefully selected based on BGP protocol 
specifications and security research.
"""
        self.doc.add_paragraph(intro)
        
        features = [
            ('announcements', 'Number of BGP route announcements', 
             'Indicates new routes being advertised. Sudden spikes may indicate route hijacks or misconfigurations.'),
            ('withdrawals', 'Number of BGP route withdrawals', 
             'Indicates routes being removed. High values suggest network instability or attacks.'),
            ('total_updates', 'Total BGP updates (announcements + withdrawals)', 
             'Overall BGP activity level. Extreme values indicate churn or DDoS patterns.'),
            ('withdrawal_ratio', 'Ratio of withdrawals to announcements', 
             'Normalized metric (0-1). Values >0.7 indicate excessive route instability.'),
            ('flap_count', 'Number of route flapping events', 
             'Routes rapidly appearing and disappearing. High values indicate peer issues or route oscillation.'),
            ('path_length', 'Average AS path length', 
             'Number of autonomous systems in routing path. Abnormally long paths (>15) may indicate route leaks or loops.'),
            ('unique_peers', 'Number of unique BGP peers', 
             'Diversity of routing sources. Sudden changes indicate peer connectivity issues.'),
            ('message_rate', 'BGP messages per second', 
             'Rate of BGP protocol messages. Spikes indicate control plane stress or attacks.'),
            ('session_resets', 'BGP session reset count', 
             'Number of TCP connection resets. High values indicate peer instability or deliberate disruption.')
        ]
        
        for feature_name, short_desc, long_desc in features:
            self.doc.add_heading(f'{feature_name}', 3)
            p = self.doc.add_paragraph()
            p.add_run(f'{short_desc}\n').bold = True
            p.add_run(long_desc)
            
        self.doc.add_page_break()
        
    def add_file_documentation(self, filename, purpose, key_functions, code_snippet=None):
        """Add documentation for a specific file"""
        self.doc.add_heading(filename, 2)
        
        self.doc.add_heading('Purpose:', 3)
        self.doc.add_paragraph(purpose)
        
        if key_functions:
            self.doc.add_heading('Key Functions/Components:', 3)
            for func_name, func_desc in key_functions:
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(f'{func_name}: ').bold = True
                p.add_run(func_desc)
                
        if code_snippet:
            self.doc.add_heading('Key Code Snippet:', 3)
            code_para = self.doc.add_paragraph(code_snippet)
            code_para.style = 'Normal'
            # Make it look like code
            for run in code_para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
    
    def add_drift_detection_section(self):
        """Add comprehensive drift detection documentation"""
        self.doc.add_heading('7. Drift Detection & Self-Healing System', 1)
        
        intro = """
The system includes an advanced self-healing capability that automatically detects model 
performance degradation (drift) and triggers retraining to maintain accuracy over time. 
This ensures the system can operate autonomously for extended periods without manual intervention.
        """
        self.doc.add_paragraph(intro)
        
        self.doc.add_heading('7.1 Drift Monitoring', 2)
        
        monitoring_desc = """
services/drift_monitor.py - Continuously monitors all 3 models for performance degradation

How it works:
• Runs hourly checks comparing recent performance vs historical baseline
• Baseline Window: 7 days of historical data
• Current Window: Last 24 hours of data
• Detects drift using statistical comparison of key metrics
• Creates drift flags in database when degradation detected
• Generates detailed drift reports with metrics and timestamps
        """
        self.doc.add_paragraph(monitoring_desc)
        
        self.doc.add_heading('Drift Detection Metrics:', 3)
        drift_metrics = [
            'Isolation Forest: Anomaly rate shifts > 10%, average score changes > 0.15',
            'LSTM: Reconstruction error changes > 0.20, anomaly rate shifts > 10%',
            'Heuristic: Threshold breach rate changes > 15%, rule trigger frequency shifts',
            'All Models: Statistical significance testing (p < 0.05)'
        ]
        for metric in drift_metrics:
            self.doc.add_paragraph(metric, style='List Bullet')
        
        self.doc.add_heading('7.2 Automatic Retraining', 2)
        
        retraining_desc = """
When drift is detected, the system automatically triggers the appropriate retraining script:

retrain_isolation_forest.py (30 seconds)
• Extracts last 7 days of data (~10,000 samples)
• Trains new Isolation Forest with 200 estimators
• Validates on holdout set (20% split)
• Computes new anomaly thresholds
• Saves model with timestamp backup
• Updates active model with hot-swap

retrain_heuristic.py (10 seconds)
• Extracts last 7 days of data
• Computes 95th percentile for each feature
• Updates 6 dynamic thresholds:
  - total_updates_critical & high
  - path_length_severe
  - withdrawal_ratio_mass
  - flap_count_high
  - message_rate_spike
• Saves updated thresholds with timestamp
• Hot-swaps into production

retrain_lstm.py (5-10 minutes)
• Extracts 7-14 days of data for temporal patterns
• Creates 10-timestep sequences
• Trains LSTM autoencoder (50 epochs max)
• Early stopping on validation loss
• Computes 95th percentile reconstruction error threshold
• Saves model weights, scaler, config
• Hot-swaps into production with zero downtime
        """
        self.doc.add_paragraph(retraining_desc)
        
        self.doc.add_heading('7.3 Hot-Swap Mechanism', 2)
        
        hotswap_desc = """
All retraining scripts implement zero-downtime model updates:

Process:
1. New model is trained and saved with timestamp suffix
2. Validation metrics are computed and logged
3. Backup of current model is created: model_name.backup_YYYYMMDD_HHMMSS
4. New model is copied to primary location (overwrites old)
5. Active detectors automatically reload model on next prediction cycle
6. No service restart required - seamless transition

Benefits:
• Zero downtime during retraining
• Automatic rollback capability (backups retained)
• Version control with timestamp-based naming
• Safe deployment with validation checks
        """
        self.doc.add_paragraph(hotswap_desc)
        
        self.doc.add_heading('Testing & Validation:', 2)
        
        testing_desc = """
test_retraining.py - Comprehensive test suite for drift system

Tests performed:
1. Isolation Forest Retraining Test
   - Verifies 30-second training time
   - Validates model file creation
   - Checks anomaly rate calculations
   - Result: ✅ PASSED (1111 samples, 1.08% anomaly rate)

2. Heuristic Threshold Update Test
   - Verifies 10-second update time
   - Validates JSON threshold file
   - Checks 6 threshold values
   - Result: ✅ PASSED (1111 samples, 15.03% anomaly rate)

3. LSTM Retraining Test
   - Verifies 5-10 minute training time
   - Validates model architecture preservation
   - Checks sequence creation (10 timesteps)
   - Result: ✅ PASSED (1104 sequences, final loss 0.316)

4. Hot-Swap Mechanism Test
   - Verifies backup file creation
   - Validates timestamp naming
   - Checks file integrity
   - Result: ✅ PASSED (2 backups created)

Overall: 4/4 tests passed successfully
        """
        self.doc.add_paragraph(testing_desc)
        
        self.doc.add_page_break()
    
    def add_dashboard_section(self):
        """Add dashboard documentation"""
        self.doc.add_heading('5. Real-time Dashboard', 1)
        
        dashboard_desc = """
dashboard_clean.html - Live web dashboard for real-time monitoring

Features:
• 10-second automatic refresh for live updates
• Chart.js for smooth, interactive visualizations
• Responsive design for desktop and mobile
• Direct integration with FastAPI backend
        """
        self.doc.add_paragraph(dashboard_desc)
        
        self.doc.add_heading('Dashboard Components:', 2)
        
        components = [
            ('Stats Cards (4 cards)',
             'Total Alerts: Count of all detections\n'
             'Active Threats: CRITICAL + HIGH severity count\n'
             'Last Hour: Recent detection count\n'
             'System Status: Operational indicator'),
            
            ('Time Series Chart',
             'Line graph showing detection count over last 24 hours\n'
             'Updates every 10 seconds with new data\n'
             'Color-coded by severity level\n'
             'Hover tooltips with exact values'),
            
            ('Model Breakdown Pie Chart',
             'Shows detection distribution across 3 models\n'
             'LSTM (blue), Isolation Forest (green), Heuristic (red)\n'
             'Percentage labels for each slice\n'
             'Interactive click to highlight'),
            
            ('Recent Detections Table',
             'Last 10 detections with full details\n'
             'Columns: Timestamp, Severity, Model, ASN, Score\n'
             'Color-coded severity badges\n'
             'Auto-scrolling for readability')
        ]
        
        for comp_name, comp_desc in components:
            self.doc.add_heading(comp_name, 3)
            self.doc.add_paragraph(comp_desc)
        
        self.doc.add_heading('How to Access Dashboard:', 2)
        access_steps = """
1. Start API server:
   python api.py
   
2. Open browser and navigate to:
   http://localhost:8000/dashboard
   
3. Dashboard will auto-refresh every 10 seconds
4. All data is live from PostgreSQL database
        """
        self.doc.add_paragraph(access_steps)
        
        self.doc.add_page_break()
    
    def add_database_architecture_section(self):
        """Add database documentation"""
        self.doc.add_heading('6. Database Architecture', 1)
        
        db_intro = """
PostgreSQL 17.7 - Primary data store with auto-cleanup triggers

Database: bgp_monitor
Host: localhost:5432
Retention: 30 days
Max Records per Table: 50,000 (auto-cleanup triggers)
        """
        self.doc.add_paragraph(db_intro)
        
        self.doc.add_heading('Database Tables:', 2)
        
        tables = [
            ('bgp_raw_data',
             'Raw BGP messages from RIS Live\n'
             'Columns: id, timestamp, host, peer, peer_asn, type, path, origin, next_hop, prefixes\n'
             'Retention: 50,000 records (FIFO)\n'
             'Indexes: timestamp, peer_asn, type'),
            
            ('bgp_features',
             '9 extracted BGP features in 30-second windows\n'
             'Columns: id, window_start, window_end, announcements, withdrawals, total_updates, '
             'withdrawal_ratio, flap_count, path_length, unique_peers, message_rate, session_resets\n'
             'Retention: 50,000 records\n'
             'Indexes: window_start, window_end'),
            
            ('lstm_detections',
             'LSTM autoencoder detection results\n'
             'Columns: id, timestamp, reconstruction_error, threshold, is_anomaly, confidence\n'
             'Retention: 50,000 records\n'
             'Indexes: timestamp, is_anomaly'),
            
            ('if_detections',
             'Isolation Forest detection results\n'
             'Columns: id, timestamp, anomaly_score, threshold, is_anomaly, confidence\n'
             'Retention: 50,000 records\n'
             'Indexes: timestamp, is_anomaly'),
            
            ('heuristic_detections',
             'Heuristic rule-based detection results\n'
             'Columns: id, timestamp, triggered_rules, score, is_anomaly, reasons\n'
             'Retention: 50,000 records\n'
             'Indexes: timestamp, is_anomaly'),
            
            ('ensemble_results',
             'Combined ensemble voting results\n'
             'Columns: id, timestamp, lstm_score, if_score, heuristic_score, ensemble_score, '
             'severity, is_anomaly\n'
             'Retention: 50,000 records\n'
             'Indexes: timestamp, severity, is_anomaly'),
            
            ('correlated_incidents',
             'Grouped related detections\n'
             'Columns: id, incident_start, incident_end, detection_count, severity, asn_list, '
             'prefix_list, description\n'
             'Retention: 30 days\n'
             'Indexes: incident_start, severity'),
            
            ('drift_flags',
             'Active drift detection flags\n'
             'Columns: id, model_name, flag_type, created_at, details\n'
             'Retention: 90 days\n'
             'Indexes: model_name, created_at, flag_type'),
            
            ('drift_reports',
             'Historical drift detection reports\n'
             'Columns: id, model_name, report_time, baseline_start, baseline_end, current_start, '
             'current_end, metrics, drift_detected, recommendations\n'
             'Retention: 180 days\n'
             'Indexes: model_name, report_time')
        ]
        
        for table_name, table_desc in tables:
            self.doc.add_heading(table_name, 3)
            self.doc.add_paragraph(table_desc)
        
        self.doc.add_heading('Auto-Cleanup Triggers:', 2)
        
        triggers_desc = """
Database triggers automatically maintain 50,000 record limits:

Trigger Logic:
1. After each INSERT, check table row count
2. If count > 50,000, delete oldest records
3. Keep newest 50,000 records only
4. Uses timestamp/id for ordering

Benefits:
• Automatic space management
• No manual cleanup required
• Consistent performance (queries stay fast)
• Prevents database bloat
• 30-day retention at current data rates

Example Trigger:
CREATE OR REPLACE FUNCTION cleanup_bgp_raw_data()
RETURNS TRIGGER AS $$
BEGIN
    DELETE FROM bgp_raw_data
    WHERE id IN (
        SELECT id FROM bgp_raw_data
        ORDER BY timestamp DESC
        OFFSET 50000
    );
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER bgp_raw_data_cleanup
AFTER INSERT ON bgp_raw_data
FOR EACH STATEMENT
EXECUTE FUNCTION cleanup_bgp_raw_data();
        """
        self.doc.add_paragraph(triggers_desc)
        
        self.doc.add_page_break()
    
    def add_complete_restart_commands(self):
        """Add comprehensive restart and usage guide"""
        self.doc.add_heading('11. Usage Guide - Complete Restart Commands', 1)
        
        intro = """
This section provides ALL commands needed to restart the system from scratch.
Follow these steps in order after system reboot or maintenance.
        """
        self.doc.add_paragraph(intro)
        
        self.doc.add_heading('Step 1: Verify Prerequisites', 2)
        prereq_check = """
# Check Python version (need 3.8+)
python --version

# Check PostgreSQL status
psql --version
psql -U postgres -d bgp_monitor -c "SELECT version();"

# Check virtual environment exists
ls venv/  # Should show: Scripts/, Lib/, pyvenv.cfg
        """
        code_para = self.doc.add_paragraph(prereq_check)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_heading('Step 2: Activate Virtual Environment', 2)
        venv_cmd = """
# Windows (PowerShell)
.\\venv\\Scripts\\Activate.ps1

# Windows (Command Prompt)
.\\venv\\Scripts\\activate.bat

# Linux/Mac
source venv/bin/activate

# Verify activation (should show venv path)
which python
        """
        code_para = self.doc.add_paragraph(venv_cmd)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_heading('Step 3: Check Database Connection', 2)
        db_check = """
# Test database connection
python -c "import psycopg2; conn = psycopg2.connect(
    dbname='bgp_monitor', user='postgres', password='your_password',
    host='localhost', port='5432'); print('✓ Database connected'); conn.close()"

# Check table counts
psql -U postgres -d bgp_monitor -c "
SELECT 
    'bgp_raw_data' as table_name, COUNT(*) as record_count FROM bgp_raw_data
UNION ALL SELECT 'bgp_features', COUNT(*) FROM bgp_features
UNION ALL SELECT 'lstm_detections', COUNT(*) FROM lstm_detections
UNION ALL SELECT 'if_detections', COUNT(*) FROM if_detections
UNION ALL SELECT 'heuristic_detections', COUNT(*) FROM heuristic_detections
UNION ALL SELECT 'ensemble_results', COUNT(*) FROM ensemble_results
UNION ALL SELECT 'correlated_incidents', COUNT(*) FROM correlated_incidents;
"
        """
        code_para = self.doc.add_paragraph(db_check)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_heading('Step 4: Start All 8 Services', 2)
        services_cmd = """
# Start all services with one command
python run_all_services.py

# This starts:
# 1. RIS Live Collector (WebSocket data collection)
# 2. Feature Aggregator (30-second feature extraction)
# 3. LSTM Detector (temporal pattern detection)
# 4. Isolation Forest Detector (statistical outlier detection)
# 5. Heuristic Detector (rule-based detection)
# 6. Ensemble Coordinator (weighted voting)
# 7. Correlation Engine (incident grouping)
# 8. Drift Monitor (self-healing system)

# Services will run in background with PID files
# Logs written to: *.log files in project directory

# Check service status
python check_drift_status.py  # Shows all service status
        """
        code_para = self.doc.add_paragraph(services_cmd)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_heading('Step 5: Start API Server', 2)
        api_cmd = """
# Start FastAPI server (port 8000)
python api.py

# API will be available at:
# - Dashboard: http://localhost:8000/dashboard
# - API Docs: http://localhost:8000/docs
# - Health Check: http://localhost:8000/health

# Test API endpoints
curl http://localhost:8000/api/live-stats
curl http://localhost:8000/api/time-series?hours=24
curl http://localhost:8000/api/recent-detections?limit=10
        """
        code_para = self.doc.add_paragraph(api_cmd)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_heading('Step 6: Monitor System Health', 2)
        monitor_cmd = """
# Check real-time data freshness
python check_realtime.py

# Expected output:
# ✓ Recent data found: 5.2 minutes old
# ✓ Data is reasonably fresh (target: < 10 min)

# Check drift status
python check_drift_status.py

# Expected output:
# Drift Monitor: Active (0.5 minutes ago)
# LSTM: No drift
# Isolation Forest: No drift
# Heuristic: No drift
# Last Retraining: 2 hours ago

# View recent alerts
python check_alerts.py
        """
        code_para = self.doc.add_paragraph(monitor_cmd)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_heading('Step 7: Stop All Services (When Needed)', 2)
        stop_cmd = """
# Stop all services gracefully
# Method 1: Using PID files
cat .ris_collector.pid | xargs kill
cat .feature_aggregator.pid | xargs kill
cat .lstm_detector.pid | xargs kill
cat .if_detector.pid | xargs kill
cat .heuristic_detector.pid | xargs kill
cat .ensemble_coordinator.pid | xargs kill
cat .correlation_engine.pid | xargs kill
cat .drift_monitor.pid | xargs kill

# Method 2: Kill by process name (Windows)
taskkill /F /IM python.exe /FI "WINDOWTITLE eq *services*"

# Method 3: Kill by process name (Linux)
pkill -f "services/"

# Stop API server
# Press Ctrl+C in the terminal where api.py is running
        """
        code_para = self.doc.add_paragraph(stop_cmd)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_heading('Troubleshooting Common Issues:', 2)
        
        troubleshooting = [
            ('Issue: "Module not found" errors',
             'Solution: Ensure virtual environment is activated\n'
             'Command: .\\venv\\Scripts\\Activate.ps1\n'
             'Verify: which python (should show venv path)'),
            
            ('Issue: Database connection refused',
             'Solution: Check PostgreSQL is running\n'
             'Windows: services.msc → Find PostgreSQL → Start\n'
             'Linux: sudo systemctl start postgresql'),
            
            ('Issue: Port 8000 already in use',
             'Solution: Kill existing process using port\n'
             'Windows: netstat -ano | findstr :8000 → taskkill /PID <PID> /F\n'
             'Linux: lsof -i :8000 → kill <PID>'),
            
            ('Issue: Services not starting',
             'Solution: Check for stale PID files\n'
             'Command: rm .*.pid\n'
             'Then restart: python run_all_services.py'),
            
            ('Issue: No data in database',
             'Solution: Check RIS Live collector log\n'
             'Command: tail -f ris_collector.log\n'
             'Should show "Stored X messages" every 5 seconds'),
            
            ('Issue: Dashboard shows old data',
             'Solution: Restart all services\n'
             'Services should process data continuously\n'
             'Check logs for errors: tail -f *.log'),
            
            ('Issue: Models not loading',
             'Solution: Verify model files exist\n'
             'Check: ls model_output/lstm/lstm_best.h5\n'
             'Check: ls model_output/isolation_forest.pkl\n'
             'Check: ls model_output/heuristic_thresholds.json\n'
             'If missing: Retrain models using Section 10')
        ]
        
        for issue_title, issue_solution in troubleshooting:
            self.doc.add_heading(issue_title, 3)
            self.doc.add_paragraph(issue_solution)
        
        self.doc.add_page_break()
    
    def add_training_commands(self):
        """Add complete training commands section"""
        self.doc.add_heading('10. Model Training Commands', 1)
        
        training_intro = """
If models need to be retrained from scratch (e.g., after data corruption or 
major system changes), use these commands. Note: Automatic retraining handles 
routine updates, so manual training is rarely needed.
        """
        self.doc.add_paragraph(training_intro)
        
        self.doc.add_heading('Option 1: Automatic Retraining (Recommended)', 2)
        auto_training = """
# Let drift monitor handle retraining automatically
# Just run the test to trigger immediate retraining
python test_retraining.py

# This will:
# 1. Test and retrain Isolation Forest (30 seconds)
# 2. Test and retrain Heuristic thresholds (10 seconds)
# 3. Test and retrain LSTM model (5-10 minutes)
# 4. Verify hot-swap mechanism
# 5. All models updated with backups created
        """
        code_para = self.doc.add_paragraph(auto_training)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_heading('Option 2: Manual Individual Model Training', 2)
        manual_training = """
# Retrain Isolation Forest only (30 seconds)
python retrain_isolation_forest.py

# Output:
# ✓ Extracted 10,000 samples from last 7 days
# ✓ Trained Isolation Forest with 200 estimators
# ✓ Anomaly rate: 1.08%
# ✓ Model saved: model_output/isolation_forest.pkl
# ✓ Backup created: model_output/isolation_forest.backup_20260108_181500

# Retrain Heuristic thresholds only (10 seconds)
python retrain_heuristic.py

# Output:
# ✓ Extracted 1,111 samples
# ✓ Updated 6 thresholds at 95th percentile
# ✓ Anomaly rate: 15.03%
# ✓ Thresholds saved: model_output/heuristic_thresholds.json
# ✓ Backup created: model_output/heuristic_thresholds.backup_20260108_181520

# Retrain LSTM model only (5-10 minutes)
python retrain_lstm.py

# Output:
# ✓ Extracted data from last 7-14 days
# ✓ Created 1,104 sequences (10 timesteps each)
# ✓ Training with early stopping (max 50 epochs)
# Epoch 1/50 - loss: 0.532
# Epoch 2/50 - loss: 0.421
# ...
# Epoch 15/50 - loss: 0.316 (early stopped)
# ✓ Final loss: 0.316
# ✓ Model saved: model_output/lstm/lstm_best.h5
# ✓ Scaler saved: model_output/lstm/scaler.pkl
# ✓ Config saved: model_output/lstm/config.json
# ✓ Backup created: model_output/lstm/lstm_best.backup_20260108_182200
        """
        code_para = self.doc.add_paragraph(manual_training)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_heading('Option 3: Full System Training from Historical Data', 2)
        full_training = """
# If you have historical CSV data for initial training
# (This was used for original system setup)

# Train LSTM from CSV
python run_training.py
# Expects: data/synthetic_30d.csv
# Creates: model_output/lstm/* (all LSTM artifacts)

# Train Isolation Forest from same data
python train_if_model.py
# Uses: data/synthetic_30d.csv
# Creates: model_output/isolation_forest.pkl

# Heuristic thresholds are computed automatically
# during first detection run based on data distribution
        """
        code_para = self.doc.add_paragraph(full_training)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_heading('Verify Training Results:', 2)
        verify_training = """
# Check model files exist
ls -lh model_output/

# Expected output:
# isolation_forest.pkl (5.78 MB)
# heuristic_thresholds.json (0.98 KB)
# lstm/
#   lstm_best.h5 (647 KB)
#   scaler.pkl (2 KB)
#   config.json (1 KB)
#   evaluation.json (1 KB)

# Test models with sample data
python test_system.py

# Expected output:
# ✓ LSTM model loaded successfully
# ✓ Isolation Forest model loaded successfully
# ✓ Heuristic thresholds loaded successfully
# ✓ All models operational
        """
        code_para = self.doc.add_paragraph(verify_training)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        self.doc.add_page_break()
                
    def generate_complete_documentation(self):
        """Generate the complete documentation"""
        print("🚀 Starting documentation generation...")
        
        # Title and TOC
        self.add_title_page()
        self.add_table_of_contents()
        
        # Main sections
        self.add_executive_summary()
        self.add_architecture_overview()
        self.add_workflow_diagram()
        
        # Core components section
        self.doc.add_heading('4. Core Components', 1)
        self.add_core_components_overview()
        self.doc.add_page_break()
        
        # NEW: Dashboard section (Section 5)
        self.add_dashboard_section()
        
        # NEW: Database architecture (Section 6)
        self.add_database_architecture_section()
        
        # NEW: Drift detection system (Section 7)
        self.add_drift_detection_section()
        
        # File-by-file documentation (Section 8)
        self.doc.add_heading('8. File-by-File Documentation', 1)
        
        # Add documentation for each major file
        self.add_data_generation_files()
        self.add_model_training_files()
        self.add_detection_files()
        self.add_utility_files()
        
        # Additional sections
        self.add_bgp_features_section()  # Section 9
        self.add_training_commands()      # Section 10
        self.add_complete_restart_commands()  # Section 11
        self.add_configuration_section()  # Section 12
        self.add_installation_guide()     # Section 10 (Installation)
        self.add_usage_guide()            # Section 8 (Usage)
        
        # Save document
        print(f"💾 Saving documentation to {self.output_path}...")
        self.doc.save(self.output_path)
        print(f"✅ Documentation generated successfully!")
        print(f"📄 File location: {os.path.abspath(self.output_path)}")
        print(f"\n📊 Document includes:")
        print(f"  ✓ Executive Summary with system status")
        print(f"  ✓ 8-Service Microservices Architecture")
        print(f"  ✓ Complete workflow diagrams")
        print(f"  ✓ Dashboard documentation")
        print(f"  ✓ Database architecture with auto-cleanup")
        print(f"  ✓ Drift detection & self-healing system")
        print(f"  ✓ Complete file-by-file documentation")
        print(f"  ✓ All restart commands")
        print(f"  ✓ Training commands for all 3 models")
        print(f"  ✓ Troubleshooting guide")
        print(f"\n🎉 Ready to explain and restart the system anytime!")
        
    def add_data_generation_files(self):
        """Document data generation related files"""
        self.doc.add_heading('5.1 Data Generation & Collection', 2)
        
        # bmp_generator.py
        self.add_file_documentation(
            'bmp_generator.py',
            """
Simulates BGP BMP (BGP Monitoring Protocol) messages according to RFC 7854. 
This module generates realistic BGP traffic patterns including normal behavior 
and various anomaly types. It supports multiple routers, peers, and complete 
BGP attribute types.
            """,
            [
                ('MultiRouterSimulator', 'Main simulator class that orchestrates multiple router instances'),
                ('RouterConfig', 'Configuration dataclass for router parameters (ID, AS number, peers)'),
                ('PeerConfig', 'Configuration for BGP peer connections'),
                ('generate_structured_data()', 'Creates time-series BGP data with specified characteristics'),
                ('generate_bmp_initiation()', 'Creates BMP initiation messages'),
                ('generate_peer_up()', 'Simulates BGP peer establishment'),
                ('generate_route_monitoring()', 'Creates BGP UPDATE messages wrapped in BMP'),
                ('generate_statistics_report()', 'Produces BMP statistics messages')
            ],
            """
# Example: Initialize simulator
peer1 = PeerConfig(peer_ip="192.168.1.10", peer_as=65001, peer_bgp_id="10.0.0.1")
router = RouterConfig(router_id="172.16.0.1", router_as=64500, 
                      sys_name="Core-Router-1", peers=[peer1])
sim = MultiRouterSimulator("127.0.0.1", 9999)
sim.add_router(router)
data_list = sim.generate_structured_data(count=1000)
            """
        )
        
        # stream_generator.py
        self.add_file_documentation(
            'stream_generator.py',
            """
Kafka producer that streams BGP data to the message queue. Reads preprocessed 
BGP datasets and publishes them to the 'bgp-stream' topic for real-time consumption.
Supports controlled data rates for testing and demonstration.
            """,
            [
                ('KafkaProducer', 'Kafka client configured for localhost:9092'),
                ('value_serializer', 'JSON serialization for message payload'),
                ('Streaming loop', 'Reads CSV and sends records with configurable delay')
            ],
            """
import pandas as pd
from kafka import KafkaProducer
import json
import time

producer = KafkaProducer(
    bootstrap_servers=['localhost:9092'],
    value_serializer=lambda x: json.dumps(x).encode('utf-8')
)

df = pd.read_csv('processed_distributed_test_nimda.csv')
for _, row in df.iterrows():
    producer.send('bgp-stream', value=row.to_dict())
    time.sleep(1)  # 1 record per second
            """
        )
        
        # db_connector.py
        self.add_file_documentation(
            'db_connector.py',
            """
Database abstraction layer for PostgreSQL operations. Handles connection management, 
data insertion, and querying for both raw BGP data and anomaly detection results.
Uses environment variables for secure credential management.
            """,
            [
                ('DBConnector.__init__()', 'Loads credentials from .env file'),
                ('connect()', 'Establishes PostgreSQL connection with error handling'),
                ('insert_raw_data()', 'Bulk insert of BGP messages'),
                ('fetch_data_for_ml()', 'Retrieves time-windowed data for ML processing'),
                ('insert_anomaly_results()', 'Persists detection results and alerts'),
                ('close()', 'Graceful connection cleanup')
            ],
            """
class DBConnector:
    def __init__(self):
        load_dotenv()
        self.DB_NAME = os.getenv("DB_NAME")
        self.DB_USER = os.getenv("DB_USER")
        self.DB_PASSWORD = os.getenv("DB_PASSWORD")
        self.DB_HOST = os.getenv("DB_HOST", "localhost")
        self.DB_PORT = os.getenv("DB_PORT", "5432")
        self.conn = None
        
    def connect(self):
        try:
            self.conn = psycopg2.connect(
                dbname=self.DB_NAME, user=self.DB_USER,
                password=self.DB_PASSWORD, host=self.DB_HOST,
                port=self.DB_PORT
            )
            return None  # Success
        except Exception as e:
            return str(e)  # Return error message
            """
        )
        
        self.doc.add_page_break()
        
    def add_model_training_files(self):
        """Document model training files"""
        self.doc.add_heading('5.2 Model Training', 2)
        
        # bgp_lstm_pipeline.py
        self.add_file_documentation(
            'bgp_lstm_pipeline.py',
            """
Complete end-to-end pipeline for training the LSTM Autoencoder model. Implements 
unsupervised learning approach where the model learns to reconstruct normal BGP 
patterns. Anomalies are detected as inputs that produce high reconstruction error.
Includes data preprocessing, model architecture, training, evaluation, and artifact saving.
            """,
            [
                ('load_csv_to_df()', 'Loads and parses CSV with timestamp handling'),
                ('extract_bgp_features()', 'Derives all 9 BGP features with fallback logic'),
                ('create_sliding_window_sequences()', 'Converts time series to overlapping sequences'),
                ('temporal_train_test_split()', 'Splits data maintaining temporal order'),
                ('filter_normal_data()', 'Removes outliers at 95th percentile for training'),
                ('build_lstm_autoencoder()', 'Constructs encoder-decoder LSTM architecture'),
                ('train_pipeline_from_csv()', 'Main entry point orchestrating full pipeline'),
                ('compute_anomaly_threshold()', 'Calculates detection threshold from training errors'),
                ('save_artifacts()', 'Persists model, scaler, config as JSON/H5/PKL')
            ],
            """
def build_lstm_autoencoder(input_shape, encoding_dim=32):
    # Encoder
    encoder_input = layers.Input(shape=input_shape)
    x = layers.LSTM(64, activation='tanh', return_sequences=True)(encoder_input)
    x = layers.Dropout(0.2)(x)
    x = layers.LSTM(encoding_dim, activation='tanh', return_sequences=False)(x)
    
    # Decoder
    x = layers.RepeatVector(input_shape[0])(x)
    x = layers.LSTM(encoding_dim, activation='tanh', return_sequences=True)(x)
    x = layers.Dropout(0.2)(x)
    x = layers.LSTM(64, activation='tanh', return_sequences=True)(x)
    decoder_output = layers.TimeDistributed(layers.Dense(input_shape[1]))(x)
    
    autoencoder = models.Model(encoder_input, decoder_output)
    autoencoder.compile(optimizer='adam', loss='mse')
    return autoencoder
            """
        )
        
        # train_if_model.py
        self.add_file_documentation(
            'train_if_model.py',
            """
Trains the Isolation Forest model using the same preprocessing pipeline as LSTM.
Isolation Forest is an ensemble of decision trees that isolates outliers by 
randomly partitioning the feature space. Anomalies require fewer splits to isolate,
resulting in shorter average path lengths in the trees.
            """,
            [
                ('train_and_save_isolation_forest()', 'Main training function'),
                ('IsolationForest configuration', 'n_estimators=200, contamination=0.01'),
                ('Feature extraction', 'Uses same 9 features as LSTM pipeline'),
                ('Model persistence', 'Saves model and feature list using joblib')
            ],
            """
from sklearn.ensemble import IsolationForest

def train_and_save_isolation_forest(csv_path, output_dir, contamination=0.01):
    # Load and preprocess
    df = load_csv_to_df(csv_path)
    df_feat = extract_bgp_features(df)
    train_df, _ = temporal_train_test_split(df_feat, train_ratio=0.8)
    train_df_normal = filter_normal_data(train_df, percentile=95)
    
    X_train = train_df_normal[FEATURES].values
    
    # Train Isolation Forest
    model_if = IsolationForest(
        n_estimators=200,
        contamination=contamination,
        random_state=42,
        n_jobs=-1
    )
    model_if.fit(X_train)
    
    # Save artifacts
    joblib.dump(model_if, os.path.join(output_dir, 'isolation_forest.pkl'))
    joblib.dump(FEATURES, os.path.join(output_dir, 'feature_list.pkl'))
            """
        )
        
        # run_training.py
        self.add_file_documentation(
            'run_training.py',
            """
Simple training script that invokes the LSTM pipeline. Sets paths and initiates 
the training process. Acts as the entry point for model training workflow.
            """,
            [
                ('train_pipeline_from_csv()', 'Calls main LSTM training function'),
                ('Configuration', 'Specifies data path and output directory')
            ],
            """
from bgp_lstm_pipeline import train_pipeline_from_csv

train_pipeline_from_csv(
    csv_path="E:\\\\Advantal_models\\\\lstm model\\\\data\\\\synthetic_30d.csv",
    output_dir="model_output",
)
            """
        )
        
        self.doc.add_page_break()
        
    def add_detection_files(self):
        """Document detection related files"""
        self.doc.add_heading('5.3 Ensemble Detection & Real-time Monitoring', 2)
        
        # ensemble_bgp_optimized.py
        self.add_file_documentation(
            'ensemble_bgp_optimized.py',
            """
Core anomaly scoring engine that combines multiple detection methods. Implements 
weighted ensemble fusion using Z-score normalization. Scores from Isolation Forest 
and LSTM are standardized, weighted, combined, and mapped to severity levels.
Includes heuristic rules for deterministic high-confidence detections.
            """,
            [
                ('HeuristicDetector.score_row()', 'Applies rule-based thresholds'),
                ('aggregate_mrt_data()', 'Preprocesses raw MRT format data'),
                ('run_anomaly_scoring()', 'Main scoring function - loads models and scores data'),
                ('compute_ensemble_score()', 'Weighted Z-score fusion'),
                ('classify_severity()', 'Maps combined score to CRITICAL/HIGH/MEDIUM/LOW/NORMAL'),
                ('optimize_ensemble_weights()', 'Grid search for optimal IF/LSTM weights'),
                ('generate_visualizations()', 'Creates plots for analysis')
            ],
            """
class HeuristicDetector:
    @staticmethod
    def score_row(row):
        score, reasons = 0.0, []
        # Churn Rules
        if row['total_updates'] > 2000:
            score = max(score, 1.0)
            reasons.append("CRITICAL_CHURN")
        elif row['total_updates'] > 500:
            score = max(score, 0.8)
            reasons.append("HIGH_CHURN")
        # Path Rules
        if row['path_length'] > 25:
            score = max(score, 0.8)
            reasons.append("SEVERE_PATH_LEN")
        # Withdrawal Rules
        if row['withdrawal_ratio'] > 0.9:
            score = max(score, 0.9)
            reasons.append("MASS_WITHDRAWAL")
        return score, reasons

def run_anomaly_scoring(df_ml):
    # Load models
    model_if = joblib.load('model_output/isolation_forest.pkl')
    model_lstm = load_model('model_output/lstm/lstm_best.h5')
    scaler = joblib.load('model_output/lstm/scaler.pkl')
    
    # Score with IF
    if_scores = -model_if.score_samples(X)
    
    # Score with LSTM
    X_seq = create_sequences(X_scaled)
    reconstructions = model_lstm.predict(X_seq)
    lstm_errors = np.mean((X_seq - reconstructions)**2, axis=(1,2))
    
    # Ensemble fusion
    z_if = (if_scores - if_scores.mean()) / if_scores.std()
    z_lstm = (lstm_errors - lstm_errors.mean()) / lstm_errors.std()
    ensemble_score = weights['if'] * z_if + weights['lstm'] * z_lstm
    
    return results_df, thresholds
            """
        )
        
        # hybrid_detector.py
        self.add_file_documentation(
            'hybrid_detector.py',
            """
Real-time BGP anomaly detector that consumes Kafka stream, performs ML inference, 
validates routes with RPKI, and saves alerts. Maintains a sliding window buffer 
for LSTM predictions and integrates with Routinator API for ROA validation.
This is the production deployment component.
            """,
            [
                ('model_lstm', 'Loaded LSTM model for real-time prediction'),
                ('KafkaConsumer', 'Subscribed to "bgp-stream" topic'),
                ('history_buffer', 'Maintains 10-record sliding window'),
                ('validate_roa()', 'Async function querying Routinator API'),
                ('process_stream()', 'Main event loop consuming and processing messages'),
                ('Feature extraction', 'Parses Kafka message to 9-feature vector'),
                ('Database persistence', 'Saves alerts to PostgreSQL using SQLAlchemy')
            ],
            """
async def validate_roa(asn: str, prefix: str) -> dict:
    try:
        asn_number = asn.replace("AS", "")
        url = f"{ROUTINATOR_URL}/api/v1/validity/{asn_number}/{prefix}"
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(url)
            if response.status_code == 200:
                data = response.json()
                state = data.get("validated_route", {}).get("validity", {}).get("state", "not-found")
                return {"status": state.lower(), "details": data}
            return {"status": "not-found", "details": {}}
    except Exception as e:
        return {"status": "error", "details": {"error": str(e)}}

async def process_stream():
    for message in consumer:
        data = message.value
        features = [float(data.get(f, 0)) for f in FEATURES]
        
        history_buffer.append(features)
        if len(history_buffer) > SEQ_LEN: history_buffer.pop(0)
        
        if len(history_buffer) == SEQ_LEN:
            input_arr = np.asarray([history_buffer], dtype=np.float32)
            pred = model_lstm.predict(input_arr, verbose=0)
            ml_anomaly = bool(pred[0][0] > 0.5)
        
        validation = await validate_roa(data['asn'], data['prefix'])
        
        if ml_anomaly or validation['status'] != 'valid':
            # Save alert to database
            save_alert(data, ml_anomaly, validation)
            """
        )
        
        # anomaly_pipeline.py
        self.add_file_documentation(
            'anomaly_pipeline.py',
            """
Orchestrates the complete end-to-end pipeline for batch processing. Generates 
synthetic data, inserts to database, fetches for ML processing, runs ensemble 
scoring, and persists results. Useful for testing and batch analysis.
            """,
            [
                ('initialize_simulator()', 'Sets up MultiRouterSimulator configuration'),
                ('run_pipeline()', 'Main orchestration function'),
                ('Data generation', 'Creates specified number of BGP records'),
                ('Database operations', 'Insert raw data, fetch for ML, save results'),
                ('ML processing', 'Invokes ensemble scoring on fetched data')
            ],
            """
def run_pipeline():
    db = DBConnector()
    
    # 1. Generate data
    sim = initialize_simulator()
    raw_data_list = sim.generate_structured_data(count=1000)
    db.insert_raw_data(raw_data_list)
    
    # 2. Fetch for ML
    end_time = datetime.now()
    start_time = end_time - pd.Timedelta(hours=24)
    df_for_ml = db.fetch_data_for_ml(start_time, end_time)
    
    # 3. Score anomalies
    results_df, thresholds = run_anomaly_scoring(df_for_ml)
    
    # 4. Save results
    db.insert_anomaly_results(results_df)
            """
        )
        
        self.doc.add_page_break()
        
    def add_utility_files(self):
        """Document utility files"""
        self.doc.add_heading('5.4 Utility & Evaluation Files', 2)
        
        # evaluate_anomaly_detection.py
        self.add_file_documentation(
            'evaluate_anomaly_detection.py',
            """
Evaluation script for assessing detection performance. Computes metrics like 
precision, recall, F1-score, and confusion matrix when ground truth labels 
are available. Useful for model validation and threshold tuning.
            """,
            [
                ('load_results()', 'Reads detection results from CSV'),
                ('compute_metrics()', 'Calculates performance metrics'),
                ('generate_confusion_matrix()', 'Creates visualization of TP/FP/TN/FN'),
                ('threshold_analysis()', 'Evaluates performance across threshold range')
            ]
        )
        
        # visualize_results.py
        self.add_file_documentation(
            'visualize_results.py',
            """
Generates comprehensive visualizations of detection results. Creates time series 
plots, score distributions, severity breakdowns, and correlation heatmaps for 
analysis and reporting purposes.
            """,
            [
                ('plot_timeseries()', 'Anomaly scores over time'),
                ('plot_severity_distribution()', 'Bar chart of severity counts'),
                ('plot_score_histogram()', 'Distribution of ensemble scores'),
                ('plot_feature_correlations()', 'Heatmap of feature relationships')
            ]
        )
        
        # datasetfilter.py
        self.add_file_documentation(
            'datasetfilter.py',
            """
Data preprocessing utility for cleaning and filtering BGP datasets. Handles 
missing values, outlier removal, temporal filtering, and data quality checks.
            """,
            [
                ('filter_by_timerange()', 'Extracts specific temporal window'),
                ('remove_outliers()', 'Statistical outlier detection and removal'),
                ('handle_missing_values()', 'Imputation strategies for incomplete data'),
                ('validate_features()', 'Ensures all required features are present')
            ]
        )
        
        self.doc.add_page_break()
        
    def add_configuration_section(self):
        """Document configuration files"""
        self.doc.add_heading('9. Configuration Files', 1)
        
        self.doc.add_heading('ensemble_config_optimized.json', 2)
        config_desc = """
Stores optimized ensemble configuration including model weights, severity thresholds, 
and performance metrics. This file is generated during ensemble optimization and 
used by the detection pipeline.
        """
        self.doc.add_paragraph(config_desc)
        
        config_structure = """
{
  "version": "3.0_optimized_manual",
  "weights": {
    "if": 0.5,        // Isolation Forest weight
    "lstm": 0.5       // LSTM Autoencoder weight
  },
  "thresholds": {
    "critical": 2.54,  // Z-score threshold for CRITICAL alerts
    "high": 2.14,      // Z-score threshold for HIGH alerts
    "medium": 1.82,    // Z-score threshold for MEDIUM alerts
    "low": 1.58        // Z-score threshold for LOW alerts
  },
  "metrics": {
    "total_samples": 32391,
    "critical_count": 502,
    "model_agreement_count": 237
  }
}
        """
        code_para = self.doc.add_paragraph(config_structure)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            
        self.doc.add_heading('requirements.txt', 2)
        req_desc = """
Python package dependencies required for the project.
        """
        self.doc.add_paragraph(req_desc)
        
        requirements_list = [
            'psycopg2-binary - PostgreSQL database adapter',
            'python-dotenv - Environment variable management',
            'pandas - Data manipulation and analysis',
            'numpy - Numerical computing',
            'tensorflow - Deep learning framework for LSTM',
            'scikit-learn - Machine learning library for Isolation Forest',
            'joblib - Model serialization',
            'matplotlib - Plotting and visualization',
            'kafka-python - Kafka client for streaming',
            'httpx - Async HTTP client for RPKI validation'
        ]
        
        for req in requirements_list:
            self.doc.add_paragraph(req, style='List Bullet')
            
        self.doc.add_page_break()
        
    def add_installation_guide(self):
        """Add installation instructions"""
        self.doc.add_heading('7. Installation & Setup', 1)
        
        self.doc.add_heading('Prerequisites', 2)
        prereq = [
            'Python 3.8 or higher',
            'PostgreSQL 12 or higher',
            'Apache Kafka 2.13-3.9.0 or higher',
            'Routinator (for RPKI validation)',
            'Minimum 8GB RAM (16GB recommended for training)',
            'GPU optional but recommended for faster LSTM training'
        ]
        for item in prereq:
            self.doc.add_paragraph(item, style='List Bullet')
            
        self.doc.add_heading('Step-by-Step Installation', 2)
        
        steps = """
1. Clone or download the project directory
2. Create Python virtual environment:
   python -m venv venv
   source venv/bin/activate  # On Windows: venv\\Scripts\\activate
   
3. Install Python dependencies:
   pip install -r requirements.txt
   
4. Setup PostgreSQL database:
   - Create database: CREATE DATABASE bgp_monitor;
   - Run schema: psql -d bgp_monitor -f schema_functions.sql
   
5. Configure environment variables (.env file):
   DB_NAME=bgp_monitor
   DB_USER=postgres
   DB_PASSWORD=your_password
   DB_HOST=localhost
   DB_PORT=5432
   
6. Start Kafka:
   cd kafka_2.13-3.9.0
   bin/zookeeper-server-start.sh config/zookeeper.properties
   bin/kafka-server-start.sh config/server.properties
   
7. Create Kafka topic:
   bin/kafka-topics.sh --create --topic bgp-stream --bootstrap-server localhost:9092
   
8. Install and start Routinator:
   routinator init
   routinator server --http 127.0.0.1:3323
        """
        self.doc.add_paragraph(steps)
        
        self.doc.add_page_break()
        
    def add_usage_guide(self):
        """Add usage instructions"""
        self.doc.add_heading('8. Usage Guide', 1)
        
        self.doc.add_heading('Training Models', 2)
        training = """
Before detection, models must be trained on historical BGP data:

1. Prepare training data (CSV with BGP features)
2. Train LSTM Autoencoder:
   python run_training.py
   
3. Train Isolation Forest:
   python train_if_model.py
   
Models will be saved to model_output/ directory.
        """
        self.doc.add_paragraph(training)
        
        self.doc.add_heading('Running Real-time Detection', 2)
        detection = """
1. Start the Kafka stream (in one terminal):
   python stream_generator.py
   
2. Start the hybrid detector (in another terminal):
   python hybrid_detector.py
   
The detector will consume messages and save alerts to the database.
        """
        self.doc.add_paragraph(detection)
        
        self.doc.add_heading('Batch Processing', 2)
        batch = """
For batch analysis of historical data:
   python anomaly_pipeline.py
   
This will generate data, process it, and save results.
        """
        self.doc.add_paragraph(batch)
        
        self.doc.add_heading('Viewing Results', 2)
        results = """
Query the PostgreSQL database to view detection results:
   SELECT * FROM anomaly_results WHERE severity IN ('CRITICAL', 'HIGH');
   
Or use the visualization script:
   python visualize_results.py
        """
        self.doc.add_paragraph(results)
        
        self.doc.add_page_break()


def main():
    """Main execution function"""
    print("=" * 60)
    print("BGP Anomaly Detection - Documentation Generator")
    print("=" * 60)
    
    generator = ProjectDocumentationGenerator(
        output_path="BGP_Anomaly_Detection_Complete_Documentation.docx"
    )
    
    try:
        generator.generate_complete_documentation()
        print("\n" + "=" * 60)
        print("SUCCESS! Documentation is ready.")
        print("=" * 60)
        
        # Check if conversion to PDF is requested
        print("\n📌 To convert to PDF:")
        print("   Option 1: Open the Word file and 'Save As' > PDF")
        print("   Option 2: Use online converter (e.g., doc2pdf.com)")
        print("   Option 3: Install python-docx2pdf and run:")
        print("            pip install docx2pdf")
        print("            python -c \"from docx2pdf import convert; convert('BGP_Anomaly_Detection_Complete_Documentation.docx')\"")
        
    except Exception as e:
        print(f"\n❌ Error generating documentation: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
